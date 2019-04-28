from django.shortcuts import render
from django.views import View
from django.http import JsonResponse
from django.template.loader import render_to_string

from django.core.mail import EmailMessage
from django.conf import settings
import os

from django.contrib.auth.models import User

from .models import NoticeTemplate
from .models import Office
from .models import AdditionalValue

from .forms import NoticeForm

from django.utils import timezone
import re
import time
from docx import Document


date_for_form = timezone.localtime(timezone.now()).strftime("%d.%m.%Y")


def change_file_notice(dict_for_change_file_notice, path_file_notice):
    date_for_file_save = timezone.localtime(timezone.now()).strftime("%d.%m.%Y_%H%M%S")

    dict_f_n = dict_for_change_file_notice
    path = path_file_notice

    def docx_replace_regex(doc_obj, regex, replace):
        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_replace_regex(cell, regex, replace)

    doc = Document(path)

    for key, value in dict_f_n.items():
        print(key + ' -> ' + value)

        key = re.compile(key)
        docx_replace_regex(doc, key, value)

    folder_for_send_notice = os.path.join(settings.MEDIA_ROOT, 'notice_forms/send_notices/')
    path_send_notice = folder_for_send_notice + dict_f_n['jk_object'] + '_' + dict_f_n['fioclientform'] + '_' + date_for_file_save + '.docx'
    doc.save(path_send_notice)

    return path_send_notice


def notice_create(request):
    # # имя и email пользователя
    # try:
    #     user_agent_name = request.user
    # except:
    #     user_agent_name = ''
    # try:
    #     user_agent_email = request.user.email
    # except:
    #     user_agent_email = ''

    # получаем email ипотечного специалиста
    try:
        email_for_lawyer = AdditionalValue.objects.get(name='email_type_for_lawyer').value.replace(' ', '').split(',')
    except:
        email_for_lawyer = None

    # получаем названия агентства недвижимости
    try:
        real_estate_name = AdditionalValue.objects.get(name='real_estate_agency_name').value
    except:
        real_estate_name = ''

    # текст email сообщения
    try:
        email_message = AdditionalValue.objects.get(name='email_message_text').value
    except:
        email_message = ''

    # получаем данные из шаблона ajax get-методом при первом открытии формы
    try:
        data_object_name = request.GET.get('data-object-name')
    except:
        data_object_name = ''
    try:
        data_obj_dev_name = request.GET.get('data-obj-dev-name')
    except:
        data_obj_dev_name = ''

    not_n = request.GET.get('data-notice-name')
    # print('request.GET.get -> ', not_n)
    if not_n is not None:
        try:
            notice_templ = NoticeTemplate.objects.get(nick_name=not_n)
            # print(notice_templ.fio_client)
        except:
            notice_templ = 'Привет1'
            # print(notice_templ)

    data = dict()

    if request.method == 'POST':

        not_n = request.POST['notice_templ']
        # print('request.POST.get -> ', not_n)
        if not_n is not None:
            try:
                notice_templ = NoticeTemplate.objects.get(nick_name=not_n)
                # print(notice_templ)
            except:
                notice_templ = 'Привет2'
                # print(notice_templ)

        form = NoticeForm(request.POST, request.FILES)
        if form.is_valid():

            # получаем email офиса
            try:
                o_email = form.cleaned_data['office']
                office_email = Office.objects.get(office=o_email).email
                # print(office_email)
            except:
                form.cleaned_data['office'] = 'нет офиса'
                office_email = None

            # получаем email застройщика
            try:
                d_email = form.cleaned_data['emails_developer'].replace(' ', '').split(',')
            except:
                d_email = None

            # отправляем email
            from_email = form.cleaned_data['email_agent']
            to_email = form.cleaned_data['email_agent'].replace(' ', '').split(',')

            if office_email is not None:
                to_email.append(office_email)
            if d_email is not None:
                to_email = to_email + d_email
            if email_for_lawyer is not None:
                if form.cleaned_data['type_for_lawyer'] == 'ипотека':
                    to_email = to_email + email_for_lawyer

            # print(to_email)

            # генерируем переменные для subject email
            if form.cleaned_data['type_for_notice'] == 'бронирование':
                v_bron = '[бронирование] '
            else:
                v_bron = ''
            if o_email is not None:
                v_office_email = ' офиса ' + str(o_email)
            else:
                v_office_email = ''
            if real_estate_name != '':
                v_real_estate_name = ' (' + real_estate_name + '): '
            else:
                v_real_estate_name = ' : '
            v_jk_object = form.cleaned_data['jk_object']
            v_fio_client = form.cleaned_data['fio_client']

            subject = v_bron + 'уведомление' + v_office_email + v_real_estate_name + v_jk_object + '_' + v_fio_client

            message = form.cleaned_data['email_message']

            ###
            path_file_notice = notice_templ.file_template  # путь до файла шаблона

            # словарь для замены в уведомлениях
            dict_for_change_file_notice = {

                'jk_object': v_jk_object,
                notice_templ.fio_client: v_fio_client,
                notice_templ.phone_client: form.cleaned_data['phone_client'],
                notice_templ.fio_agent: form.cleaned_data['fio_agent'],
                notice_templ.phone_agent: form.cleaned_data['phone_agent']

            }

            path_send_notice = change_file_notice(dict_for_change_file_notice, path_file_notice)

            try:
                attach = request.FILES['attach']
            except:
                attach = None
            # print(attach)

            mail = EmailMessage(subject, message, settings.EMAIL_HOST_USER, to_email,
                                headers={
                                    # 'From': from_email,
                                    'Reply-to': from_email})

            if attach is not None:
                mail.attach(attach.name, attach.read(), attach.content_type)

            mail.attach_file(path_send_notice)  # прицепляем к письму сгенерированный (заполненный) файл уведомления

            mail.content_subtype = 'html'

            # try:
            mail.send()

            # БД
            notice = form.save(commit=False)
            notice.save()  # сохраняем в БД

            data['form_is_valid'] = True
            data['message'] = 'Отправленно успешно!'
            # except Exception as e:
            #     print('SendMail Failed')
            #     data['message'] = str(e)

        else:
            data['form_is_valid'] = False
    else:

        form = NoticeForm()

    context = {
        'form': form,
        # 'user_agent_name': user_agent_name,
        # 'user_agent_email': user_agent_email,
        'notice_templ': notice_templ,
        'data_object_name': data_object_name,
        'data_obj_dev_name': data_obj_dev_name,
        'email_message': email_message
    }

    data['html_form'] = render_to_string('gen_notices/notice_create.html', context, request=request)
    return JsonResponse(data)
